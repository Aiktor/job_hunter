from docx import Document
from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole

def get_giga(auth_key, promt, id, title, description, time):
    giga = GigaChat(
    model="GigaChat",
    credentials=auth_key,
    scope="GIGACHAT_API_PERS",
    verify_ssl_certs=False
    )
    # print(description)

    payload = Chat(
        messages=[
            Messages(
                role=MessagesRole.SYSTEM,
                content=promt
            ),
            Messages(
                role=MessagesRole.USER,
                content=description
            ),
        ]
    )

    response = giga.chat(payload)
    content = response.choices[0].message.content

    document = Document()
    document.add_heading('Резюме', level=1)
    paragraph = document.add_paragraph(content)
    document.save(f'ai_resume/{id} {title} {time}.docx')
    return content